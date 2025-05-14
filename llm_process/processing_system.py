import sqlite3
import pdfplumber
import pytesseract
from PIL import Image
import os
import uuid
import spacy
import cv2
import numpy as np
import nltk
from nltk.tokenize import sent_tokenize
from nltk.corpus import wordnet
from sentence_transformers import SentenceTransformer, util
import faiss
import json
import logging
from prompt_toolkit import PromptSession
from prompt_toolkit.history import FileHistory
from prompt_toolkit.auto_suggest import AutoSuggestFromHistory
from prompt_toolkit.completion import FuzzyWordCompleter
import pickle
import blosc
from rank_bm25 import BM25Okapi
from textblob import TextBlob
import docx
import pandas as pd
import yaml
from multiprocessing import Pool
import psutil
import time
import re
from summa import summarizer
import easyocr
import multiprocessing
import tempfile
import hashlib
from cachetools import LRUCache
from tqdm import tqdm
from langdetect import detect
import argparse
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, session
from flask_wtf import FlaskForm
from wtforms import StringField, SubmitField, FileField, SelectField, PasswordField, TextAreaField
from wtforms.validators import DataRequired, Email, Length
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import zipfile
from celery import Celery
from flask_socketio import SocketIO, emit, join_room
from sqlalchemy import create_engine, text
from redis import Redis
from alembic import command
from alembic.config import Config
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from io import BytesIO
import base64
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from sklearn.ensemble import IsolationForest
from elasticsearch import Elasticsearch
import pkg_resources
from transformers import pipeline
import hdbscan
from tenacity import retry, stop_after_attempt, wait_exponential

# Lokituksen asetukset
logging.basicConfig(filename='system.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Lataa NLTK-data
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/wordnet')
except LookupError:
    nltk.download('punkt')
    nltk.download('wordnet')

# Lataa mallit
nlp = spacy.load("en_core_web_sm")
embedder = SentenceTransformer('all-MiniLM-L6-v2')
easyocr_reader = easyocr.Reader(['en'])
EMBEDDING_CACHE = LRUCache(maxsize=1000)
ANOMALY_MODEL = IsolationForest(contamination=0.1, random_state=42)
ANOMALY_MODEL_CACHE = {'hash': None, 'model': None}

# Elasticsearch-asetukset
es = Elasticsearch([{'host': 'elasticsearch', 'port': 9200}])
if not es.ping():
    logging.warning("Elasticsearch ei ole saatavilla, palataan SQLite-hakuun.")

# Nollapisteen luokittelija tunnisteille
classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli")
CANDIDATE_LABELS = ["positive", "negative", "technical", "informal", "formal", "question", "instruction"]

# Kevyt tietopankki
KNOWLEDGE_BASE = {
    "paris": {"type": "GPE", "description": "Ranskan pääkaupunki"},
    "france": {"type": "GPE", "description": "Maa Euroopassa"},
    "eiffel tower": {"type": "LANDMARK", "description": "Ikonic torni Pariisissa"}
}

# Flask-sovelluksen asetukset
app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', os.urandom(24).hex())
app.config['UPLOAD_FOLDER'] = 'Uploads'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100 Mt raja
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users.db'
app.config['CELERY_BROKER_URL'] = os.environ.get('CELERY_BROKER_URL', 'redis://redis:6379/0')
app.config['CELERY_RESULT_BACKEND'] = os.environ.get('CELERY_RESULT_BACKEND', 'redis://redis:6379/0')
app.config['REDIS_URL'] = os.environ.get('REDIS_URL', 'redis://redis:6379/0')
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Celery-asetukset
celery = Celery(app.name, broker=app.config['CELERY_BROKER_URL'])
celery.conf.update(app.config)

# SocketIO-asetukset
socketio = SocketIO(app)

# Redis-välimuisti
redis_client = Redis.from_url(app.config['REDIS_URL'], decode_responses=True)

# Nopeusrajoitus
limiter = Limiter(
    app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"]
)

# Kirjautumisen hallinta
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Käyttäjätietokannan alustus
def init_user_db():
    with sqlite3.connect('users.db') as conn:
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS users
                     (id TEXT PRIMARY KEY, email TEXT UNIQUE, password TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS audit_log
                     (id TEXT PRIMARY KEY, user_id TEXT, action TEXT, timestamp TEXT)''')
        conn.commit()

# Käyttäjäluokka
class User(UserMixin):
    def __init__(self, id, email, password):
        self.id = id
        self.email = email
        self.password = password

# Käyttäjän lataaminen
@login_manager.user_loader
def load_user(user_id):
    with sqlite3.connect('users.db') as conn:
        c = conn.cursor()
        c.execute("SELECT id, email, password FROM users WHERE id = ?", (user_id,))
        user = c.fetchone()
    if user:
        return User(user[0], user[1], user[2])
    return None

# Kirjaa tarkastustoiminto
def log_audit_action(user_id, action):
    audit_id = str(uuid.uuid4())
    with sqlite3.connect('users.db') as conn:
        c = conn.cursor()
        c.execute("INSERT INTO audit_log (id, user_id, action, timestamp) VALUES (?, ?, ?, ?)",
                  (audit_id, user_id, action, time.strftime("%Y-%m-%d %H:%M:%S")))
        conn.commit()

# Lataa ja validoi konfiguraatio
CONFIG = {
    'min_content_length': 10,
    'max_content_length': 1000000,
    'validation_entity_threshold': 1,
    'validation_token_threshold': 5,
    'bm25_weight': 0.4,
    'semantic_weight': 0.6,
    'max_answers': 5,
    'supported_extensions': ['.pdf', '.txt', '.md', '.png', '.jpg', '.jpeg', '.docx', '.csv', '.zip'],
    'quality_threshold': 0.7,
    'summary_ratio': 0.5,
    'use_sentiment_validation': True,
    'sentiment_threshold': 0.05,
    'custom_validation_rules': [],
    'embedding_compression_level': 5,
    'items_per_page': 20,
    'cache_ttl': 3600,
    'anomaly_threshold': 0.1
}
def validate_config(config):
    """Validoi konfiguraation arvot."""
    schema = {
        'min_content_length': (int, lambda x: x > 0),
        'max_content_length': (int, lambda x: x > 1000),
        'validation_entity_threshold': (int, lambda x: x >= 0),
        'validation_token_threshold': (int, lambda x: x >= 0),
        'bm25_weight': (float, lambda x: 0 <= x <= 1),
        'semantic_weight': (float, lambda x: 0 <= x <= 1),
        'max_answers': (int, lambda x: x > 0),
        'supported_extensions': (list, lambda x: all(isinstance(e, str) for e in x)),
        'quality_threshold': (float, lambda x: 0 <= x <= 1),
        'summary_ratio': (float, lambda x: 0 < x <= 1),
        'use_sentiment_validation': (bool, lambda x: True),
        'sentiment_threshold': (float, lambda x: 0 <= x <= 1),
        'custom_validation_rules': (list, lambda x: all(isinstance(r, dict) and 'regex' in r and 'error' in r for r in x)),
        'embedding_compression_level': (int, lambda x: 1 <= x <= 9),
        'items_per_page': (int, lambda x: x > 0),
        'cache_ttl': (int, lambda x: x > 0),
        'anomaly_threshold': (float, lambda x: 0 < x < 1)
    }
    for key, (type_, check) in schema.items():
        if key not in config:
            raise ValueError(f"Puuttuva konfiguraatioavain: {key}")
        if not isinstance(config[key], type_):
            raise ValueError(f"Väärä tyyppi avaimelle {key}: odotettu {type_}, saatu {type(config[key])}")
        if not check(config[key]):
            raise ValueError(f"Virheellinen arvo avaimelle {key}: {config[key]}")
    return config

# Lataa konfiguraatio tiedostosta
try:
    with open('config.yaml', 'r') as f:
        CONFIG.update(yaml.safe_load(f))
    CONFIG = validate_config(CONFIG)
except FileNotFoundError:
    with open('config.yaml', 'w') as f:
        yaml.dump(CONFIG, f)
except ValueError as e:
    print(f"Konfiguraatiovirhe: {str(e)}")
    exit(1)

# Tietokantayhteyden hankkiminen
def get_db_connection():
    conn = sqlite3.connect('data_store.db', timeout=10)
    conn.execute("PRAGMA journal_mode=WAL")
    return conn

# Tietokannan alustus Alembicillä
def init_db():
    alembic_cfg = Config("alembic.ini")
    command.upgrade(alembic_cfg, "head")
    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS collaboration (
                     id TEXT PRIMARY KEY,
                     entry_id TEXT,
                     user_id TEXT,
                     action TEXT,
                     timestamp TEXT)''')
        conn.commit()

# Plugin-järjestelmä tiedostojen käsittelijöille
FILE_PROCESSORS = {}
def load_plugins():
    """Lataa laajennukset tiedostojen käsittelyyn."""
    for entry_point in pkg_resources.iter_entry_points('data_processing.plugins'):
        try:
            plugin = entry_point.load()
            for ext, processor in plugin().items():
                FILE_PROCESSORS[ext] = processor
                logging.debug(f"Ladattu laajennus laajennukselle: {ext}")
        except Exception as e:
            logging.error(f"Laajennuksen {entry_point.name} lataus epäonnistui: {str(e)}")

# Rekisteröi tiedostokäsittelijä
def register_processor(file_ext):
    def decorator(func):
        FILE_PROCESSORS[file_ext] = func
        return func
    return decorator

# PDF-tiedostojen käsittely
@register_processor('.pdf')
def process_pdf(file_path):
    try:
        with pdfplumber.open(file_path) as pdf:
            text = []
            metadata = {"page_count": len(pdf.pages), "sections": [], "annotations": []}
            for page in pdf.pages:
                page_text = page.extract_text(layout=True) or ""
                text.append(page_text)
                tables = page.extract_tables()
                if tables:
                    text.append("\n[Table]\n" + "\n".join([str(t) for t in tables]))
                if hasattr(page, 'annots'):
                    metadata["annotations"].extend([ann.get('contents', '') for ann in page.annots or []])
            return "\n".join(text), metadata
    except Exception as e:
        logging.error(f"PDF-käsittelyvirhe: {str(e)}")
        return f"Virhe: {str(e)}", {}

# Tekstitiedostojen käsittely
@register_processor('.txt')
@register_processor('.md')
def process_text(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
            metadata = {"file_size": os.path.getsize(file_path)}
            return text, metadata
    except Exception as e:
        logging.error(f"Tekstin käsittelyvirhe: {str(e)}")
        return f"Virhe: {str(e)}", {}

# Docx-tiedostojen käsittely
@register_processor('.docx')
def process_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        metadata = {"paragraph_count": len(doc.paragraphs)}
        return text, metadata
    except Exception as e:
        logging.error(f"Docx-käsittelyvirhe: {str(e)}")
        return f"Virhe: {str(e)}", {}

# CSV-tiedostojen käsittely
@register_processor('.csv')
def process_csv(file_path):
    try:
        df = pd.read_csv(file_path)
        text = df.to_string()
        metadata = {"rows": len(df), "columns": len(df.columns)}
        return text, metadata
    except Exception as e:
        logging.error(f"CSV-käsittelyvirhe: {str(e)}")
        return f"Virhe: {str(e)}", {}

# Kuvan esikäsittely
def preprocess_image(file_path):
    try:
        img = cv2.imread(file_path)
        if img is None:
            raise ValueError("Kuvan lataaminen epäonnistui")
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
        contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        if contours:
            x, y, w, h = cv2.boundingRect(max(contours, key=cv2.contourArea))
            roi = thresh[y:y+h, x:x+w]
        else:
            roi = thresh
        with tempfile.NamedTemporaryFile(suffix='.png', delete=True) as temp:
            temp_path = temp.name
            cv2.imwrite(temp_path, roi)
            return temp_path
    except Exception as e:
        logging.error(f"Kuvan esikäsittelyvirhe: {str(e)}")
        return f"Virhe: {str(e)}"

# Tarkista EasyOCR:n toimivuus
def check_easyocr():
    try:
        test_path = os.path.join(os.path.dirname(__file__), 'test.png')
        if not os.path.exists(test_path):
            Image.new('RGB', (100, 100)).save(test_path)
        easyocr_reader.readtext(test_path)
        if os.path.exists(test_path):
            os.remove(test_path)
        return True
    except:
        return False

# Kuvatiedostojen käsittely
@register_processor('.png')
@register_processor('.jpg')
@register_processor('.jpeg')
def process_image(file_path):
    try:
        temp_path = preprocess_image(file_path)
        if isinstance(temp_path, str) and "Error" in temp_path:
            return temp_path, {}
        image = Image.open(temp_path)
        text = pytesseract.image_to_string(image)
        use_easyocr = check_easyocr()
        if not text.strip() and use_easyocr:
            results = easyocr_reader.readtext(temp_path)
            text = "\n".join([res[1] for res in results])
        elif not text.strip():
            logging.warning("EasyOCR epäonnistui tai ei ole saatavilla; ohitetaan varavaihtoehto.")
        metadata = {"dimensions": image.size}
        return text, metadata
    except Exception as e:
        logging.error(f"Kuvan käsittelyvirhe: {str(e)}")
        return f"Virhe: {str(e)}", {}
    finally:
        if 'temp_path' in locals() and os.path.exists(temp_path):
            os.remove(temp_path)

# Tekstin normalisointi
def normalize_text(text):
    """Normalisoi teksti poistamalla ylimääräiset välilyönnit ja tunnistamalla kieli."""
    text = re.sub(r'\s+', ' ', text.strip())
    try:
        lang = detect(text[:1000])
        if lang != 'en':
            logging.info(f"Havaittu kieli: {lang}")
    except:
        pass
    return text

# Vastauksen puhdistus
def clean_answer(text):
    """Poista päällekkäiset lauseet vastauksesta."""
    text = re.sub(r'\s+', ' ', text.strip())
    sentences = sent_tokenize(text)
    seen = set()
    cleaned = [s for s in sentences if s.lower() not in seen and not seen.add(s.lower())]
    return " ".join(cleaned)

# Sisällön validointi
def validate_content(content):
    """Validoi sisällön pituus, entiteetit ja sentimentti."""
    content = normalize_text(content)
    if not content or len(content) < CONFIG['min_content_length']:
        return False, "Sisältö on tyhjä tai liian lyhyt", 0.0
    if len(content) > CONFIG['max_content_length']:
        return False, "Sisältö ylittää maksimipituuden", 0.0
    for rule in CONFIG.get('custom_validation_rules', []):
        if not re.match(rule['regex'], content):
            return False, rule['error'], 0.0
    doc = nlp(content[:1000])
    entities = len(doc.ents)
    tokens = len([t for t in doc if not t.is_stop and t.is_alpha])
    blob = TextBlob(content[:1000])
    polarity = blob.sentiment.polarity
    quality_score = (entities / max(1, CONFIG['validation_entity_threshold']) +
                     tokens / max(1, CONFIG['validation_token_threshold']) +
                     abs(polarity)) / 3
    if entities < CONFIG['validation_entity_threshold'] or tokens < CONFIG['validation_token_threshold']:
        return False, "Sisällöstä puuttuu merkityksellisiä entiteettejä tai avainsanoja", quality_score
    if CONFIG.get('use_sentiment_validation', True) and abs(polarity) < CONFIG.get('sentiment_threshold', 0.05):
        return False, "Sisällöllä on neutraali sentimentti", quality_score
    return True, "Kelvollinen", quality_score

# Poikkeamien tunnistus
def detect_anomalies(entries):
    """Tunnista poikkeavat merkinnät laatupisteiden perusteella."""
    if not entries:
        return []
    features = [[e['quality']] for e in entries]
    data_hash = hashlib.sha256(str(features).encode()).hexdigest()
    if ANOMALY_MODEL_CACHE['hash'] != data_hash:
        model = IsolationForest(contamination=0.1, random_state=42)
        predictions = model.fit_predict(features)
        ANOMALY_MODEL_CACHE.update({'hash': data_hash, 'model': model})
    else:
        predictions = ANOMALY_MODEL_CACHE['model'].predict(features)
    return [e['id'] for e, pred in zip(entries, predictions) if pred == -1]

# Upotuksen hankkiminen
def get_embedding(sentence):
    """Hae lauseen upotus välimuistista tai laske se."""
    cache_key = f"emb:{hashlib.md5(sentence.encode()).hexdigest()}"
    cached = redis_client.get(cache_key)
    if cached:
        return pickle.loads(cached)
    emb = embedder.encode(sentence, convert_to_tensor=False)
    redis_client.setex(cache_key, CONFIG['cache_ttl'], pickle.dumps(emb))
    return emb

# Elasticsearch-yhteyden varmistus
@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def safe_es_ping():
    """Varmista Elasticsearch-yhteys uudelleenyrityksillä."""
    if not es.ping():
        raise ConnectionError("Elasticsearch ei ole saatavilla")
    return True

# Metatietojen rikastaminen
def enrich_metadata(content, metadata, entry_id):
    """Rikastuta metatietoja tunnisteilla ja liittyvillä merkinnöillä."""
    # Generoi tunnisteet
    try:
        classification = classifier(content[:1000], CANDIDATE_LABELS, multi_label=True)
        tags = [label for label, score in zip(classification['labels'], classification['scores']) if score > 0.5]
        metadata['tags'] = tags
    except Exception as e:
        logging.error(f"Tunnisteiden generointivirhe: {str(e)}")
        metadata['tags'] = []

    # Klusterointi liittyville merkinnöille
    try:
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute("SELECT entry_id, embedding FROM embeddings WHERE entry_id != ?", (entry_id,))
            embeddings = [(row[0], pickle.loads(blosc.decompress(row[1]))) for row in c.fetchall()]
        if embeddings and len(embeddings) > 5:
            emb_vectors = np.array([emb for _, emb in embeddings])
            clusterer = hdbscan.HDBSCAN(min_cluster_size=2, metric='cosine')
            labels = clusterer.fit_predict(emb_vectors)
            current_emb = get_embedding(content)
            similarities = util.cos_sim(current_emb, emb_vectors).numpy().flatten()
            related = [embeddings[i][0] for i in similarities.argsort()[-3:][::-1] if labels[i] == labels[similarities.argmax()]]
            metadata['related_entries'] = related
        else:
            metadata['related_entries'] = []
    except Exception as e:
        logging.error(f"Klusterointivirhe: {str(e)}")
        metadata['related_entries'] = []
    
    return metadata

# Tietojen tallennus
def store_data(data_type, content, source, metadata, user_id=None):
    """Tallenna data tietokantaan ja indeksoi se."""
    content = normalize_text(content)
    is_valid, validation_msg, quality_score = validate_content(content)
    entry_id = str(uuid.uuid4())
    
    # Rikastuta metatietoja
    metadata = enrich_metadata(content, metadata, entry_id)
    
    with get_db_connection() as conn:
        conn.execute("BEGIN TRANSACTION")
        try:
            c = conn.cursor()
            c.execute("INSERT INTO data_entries (id, data_type, content, source, validated, metadata, quality_score) VALUES (?, ?, ?, ?, ?, ?, ?)",
                      (entry_id, data_type, content, source, 1 if is_valid else 0, json.dumps(metadata), quality_score))
            c.execute("INSERT INTO history (id, entry_id, content, timestamp) VALUES (?, ?, ?, ?)",
                      (str(uuid.uuid4()), entry_id, content, time.strftime("%Y-%m-%d %H:%M:%S")))
            sentences = [s for s in sent_tokenize(content) if s.strip() and len(s) > 5]
            if sentences:
                embeddings = embedder.encode(sentences, batch_size=32, show_progress_bar=False)
                for sent, emb in zip(sentences, embeddings):
                    sent_id = str(uuid.uuid4())
                    compressed_emb = blosc.compress(pickle.dumps(emb), clevel=CONFIG.get('embedding_compression_level', 5))
                    c.execute("INSERT INTO embeddings (sentence_id, entry_id, sentence, embedding) VALUES (?, ?, ?, ?)",
                              (sent_id, entry_id, sent, compressed_emb))
            # Indeksoi Elasticsearchiin
            if safe_es_ping():
                es.index(index='data_entries', id=entry_id, body={
                    'data_type': data_type,
                    'content': content,
                    'source': source,
                    'metadata': metadata,
                    'quality_score': quality_score
                })
            conn.commit()
        except Exception as e:
            conn.rollback()
            logging.error(f"Tietojen tallennusvirhe: {str(e)}")
            raise e
    redis_client.delete('list:*', 'review:*')
    if user_id:
        log_audit_action(user_id, f"Tallennettu merkintä {entry_id}")
        socketio.emit('collaboration_update', {
            'entry_id': entry_id,
            'action': 'create',
            'user_id': user_id,
            'content_preview': content[:100]
        }, room=entry_id)
    return entry_id, validation_msg, quality_score

# Välimuistin tyhjennys
def clear_cache(prefix):
    """Tyhjennä välimuisti tietyllä etuliitteellä."""
    cursor = '0'
    while cursor != 0:
        cursor, keys = redis_client.scan(cursor=cursor, match=f"{prefix}*", count=100)
        if keys:
            redis_client.delete(*keys)

# Kaikkien tietojen haku
def get_all_data(page=1, per_page=20, search=None, data_type=None, validated=None, metadata_query=None):
    """Hae kaikki merkinnät tietokannasta tai Elasticsearchista."""
    cache_key = f"list:{page}:{per_page}:{search or ''}:{data_type or ''}:{validated if validated is not None else ''}:{metadata_query or ''}"
    cached = redis_client.get(cache_key)
    if cached:
        return pickle.loads(cached)
    
    offset = (page - 1) * per_page
    if safe_es_ping() and (search or metadata_query):
        es_query = {
            'query': {
                'bool': {
                    'must': []
                }
            },
            'from': offset,
            'size': per_page,
            'sort': [{'quality_score': 'desc'}]
        }
        if search:
            es_query['query']['bool']['must'].append({
                'multi_match': {
                    'query': search,
                    'fields': ['content', 'source']
                }
            })
        if data_type:
            es_query['query']['bool']['must'].append({'term': {'data_type': data_type}})
        if validated is not None:
            es_query['query']['bool']['must'].append({'term': {'validated': validated}})
        if metadata_query:
            es_query['query']['bool']['must'].append({'query_string': {'query': metadata_query}})
        try:
            result = es.search(index='data_entries', body=es_query)
            entries = [{
                'id': hit['_id'],
                'type': hit['_source']['data_type'],
                'source': hit['_source']['source'],
                'status': "Validoitu" if hit['_source'].get('validated', 0) else "Tarvitsee validointia",
                'quality': hit['_source']['quality_score'],
                'metadata': hit['_source']['metadata'],
                'content_preview': hit['_source']['content'][:100]
            } for hit in result['hits']['hits']]
            total = result['hits']['total']['value']
        except Exception as e:
            logging.error(f"Elasticsearch-hakuvirhe: {str(e)}")
            entries, total = [], 0
    else:
        query = "SELECT id, data_type, content, source, validated, metadata, quality_score FROM data_entries WHERE 1=1"
        params = []
        if search:
            query += " AND (content LIKE ? OR source LIKE ?)"
            params.extend([f'%{search}%', f'%{search}%'])
        if data_type:
            query += " AND data_type = ?"
            params.append(data_type)
        if validated is not None:
            query += " AND validated = ?"
            params.append(1 if validated else 0)
        if metadata_query:
            query += " AND metadata LIKE ?"
            params.append(f'%{metadata_query}%')
        query += " ORDER BY quality_score DESC LIMIT ? OFFSET ?"
        params.extend([per_page, offset])
        
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute(query, params)
            entries = c.fetchall()
            c.execute("SELECT COUNT(*) FROM data_entries WHERE 1=1" + 
                      (" AND (content LIKE ? OR source LIKE ?)" if search else "") +
                      (" AND data_type = ?" if data_type else "") +
                      (" AND validated = ?" if validated is not None else "") +
                      (" AND metadata LIKE ?" if metadata_query else ""),
                      [p for p in params if p not in [per_page, offset]])
            total = c.fetchone()[0]
        
        entries = [{
            'id': e[0],
            'type': e[1],
            'source': e[3],
            'status': "Validoitu" if e[4] else "Tarvitsee validointia",
            'quality': e[6],
            'metadata': json.loads(e[5]),
            'content_preview': e[2][:100]
        } for e in entries]
    
    anomalies = detect_anomalies(entries)
    for entry in entries:
        entry['is_anomaly'] = entry['id'] in anomalies
    
    result = (entries, total)
    redis_client.setex(cache_key, CONFIG['cache_ttl'], pickle.dumps(result))
    return result

# Merkintöjen listaus
def list_entries(page=1, per_page=20, search=None, data_type=None, validated=None, metadata_query=None):
    """Listaa merkinnät sivutuksen ja suodattimien kanssa."""
    entries, total = get_all_data(page, per_page, search, data_type, validated, metadata_query)
    if not entries:
        return "Ei merkintöjä löydetty.", [], 0
    result = "Tallennetut merkinnät:\n" + "\n".join(
        f"Tunnus: {e['id']}, Tyyppi: {e['type']}, Lähde: {e['source']}, Tila: {e['status']}, Laatu: {e['quality']:.2f}, Metatiedot: {e['metadata']}, Poikkeama: {e['is_anomaly']}"
        for e in entries
    )
    return result, entries, total

# Tietopankin päivitys
def update_knowledge_base():
    """Päivitä tietopankki validoitujen merkintöjen entiteeteillä."""
    global KNOWLEDGE_BASE
    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute("SELECT content FROM data_entries WHERE validated = 1")
        contents = c.fetchall()
    for content in contents:
        doc = nlp(content[0][:1000])
        for ent in doc.ents:
            if ent.label_ in ["GPE", "PERSON", "ORG", "FAC"]:
                KNOWLEDGE_BASE[ent.text.lower()] = {"type": ent.label_, "description": f"{ent.label_} sisällöstä"}

# Kysymyksen uudelleenmuotoilu
def rewrite_query(question):
    """Muotoile kysymys uudelleen sen juuren perusteella."""
    question = question.lower().strip()
    doc = nlp(question)
    if doc[0].text in ("what", "where", "who", "when"):
        for token in doc:
            if token.dep_ == "ROOT":
                return token.text + " " + " ".join(t.text for t in token.subtree if t != token)
    return question

# Kysymyksen laajentaminen
def expand_query(question):
    """Laajenna kysymys synonyymeillä."""
    words = [token.text.lower() for token in nlp(question) if token.is_alpha and not token.is_stop]
    expanded = set(words)
    for word in words:
        for syn in wordnet.synsets(word):
            for lemma in syn.lemmas():
                expanded.add(lemma.name().lower().replace('_', ' '))
    return ' '.join(expanded)

# Kysymyksen intention tunnistus
def detect_intent(question):
    """Tunnista kysymyksen intentio."""
    question_lower = question.lower()
    doc = nlp(question)
    if question_lower.startswith(("list", "what are", "name")):
        return "list"
    elif question_lower.startswith("define"):
        return "definition"
    elif any(dep in [t.dep_ for t in doc] for dep in ['prep', 'advmod']) or question_lower.startswith(("why", "how")):
        return "explanation"
    elif question_lower.startswith(("compare", "difference")):
        return "comparison"
    return "factoid"

# FAISS-indeksin hallinta
def get_dataset_hash():
    """Laske tietojoukon hash-arvo."""
    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute("SELECT id, content FROM data_entries ORDER BY id")
        data = c.fetchall()
    return hashlib.sha256(str(data).encode()).hexdigest()

# FAISS-indeksin tallennus
def save_faiss_index(index, filename='faiss_index.bin'):
    """Tallenna FAISS-indeksi tiedostoon."""
    faiss.write_index(index, filename)
    with open('faiss_index.meta', 'w') as f:
        json.dump({'dataset_hash': get_dataset_hash(), 'timestamp': time.time()}, f)

# FAISS-indeksin lataaminen
def load_faiss_index(filename='faiss_index.bin', dim=384):
    """Lataa FAISS-indeksi tai luo uusi, jos se ei ole ajan tasalla."""
    current_hash = get_dataset_hash()
    if os.path.exists(filename) and os.path.exists('faiss_index.meta'):
        with open('faiss_index.meta', 'r') as f:
            meta = json.load(f)
        if meta['dataset_hash'] == current_hash:
            return faiss.read_index(filename)
    index = faiss.IndexIVFPQ(faiss.IndexFlatIP(dim), dim, 8, 8, 8)
    index.nprobe = 8
    return index

# Kysymyksiin vastaaminen
def answer_question(question):
    """Vastaa kysymykseen käyttäen FAISS- ja BM25-hakua."""
    cache_key = f"answer:{hashlib.md5(question.encode()).hexdigest()}"
    cached = redis_client.get(cache_key)
    if cached:
        return cached.decode()

    start_time = time.time()
    entries = get_all_data()[0]
    if not entries:
        return "Ei dataa saatavilla."

    intent = detect_intent(question)
    rewritten_question = rewrite_query(question)
    expanded_question = expand_query(rewritten_question)
    question_emb = get_embedding(expanded_question)

    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute("SELECT sentence_id, entry_id, sentence, embedding FROM embeddings WHERE entry_id IN (SELECT id FROM data_entries WHERE validated = 1 AND quality_score > ?)",
                  (CONFIG['quality_threshold'],))
        embeddings = c.fetchall()

    if not embeddings:
        return "Ei indeksoituja lauseita löydetty."

    sentences = [emb[2] for emb in embeddings]
    tokenized_sentences = [[w.text for w in nlp(sent.lower()) if w.is_alpha and not w.is_stop] for sent in sentences]
    bm25 = BM25Okapi(tokenized_sentences)

    dim = len(get_embedding("test"))
    index = load_faiss_index(dim=dim)
    emb_vectors = np.array([pickle.loads(blosc.decompress(emb[3])) for emb in embeddings]).astype('float32')
    current_hash = get_dataset_hash()
    if os.path.exists('faiss_index.meta'):
        with open('faiss_index.meta', 'r') as f:
            meta = json.load(f)
        if meta['dataset_hash'] != current_hash and len(emb_vectors) > 0:
            index = faiss.IndexIVFPQ(faiss.IndexFlatIP(dim), dim, 8, 8, 8)
            index.nprobe = 8
            index.train(emb_vectors)
    if len(emb_vectors) > 0:
        index.add(emb_vectors)

    bm25_scores = bm25.get_scores([w.text for w in nlp(expanded_question.lower()) if w.is_alpha and not w.is_stop])
    k = min(CONFIG['max_answers'], len(embeddings))
    distances, indices = index.search(question_emb.reshape(1, -1).astype('float32'), k)

    combined_scores = {}
    for idx, dist in zip(indices[0], distances[0]):
        entry_id = embeddings[idx][1]
        score = CONFIG['semantic_weight'] * dist + CONFIG['bm25_weight'] * bm25_scores[idx]
        combined_scores[idx] = (score, entry_id, embeddings[idx][2])

    answers = []
    seen_entries = set()
    for idx in sorted(combined_scores, key=lambda x: combined_scores[x][0], reverse=True):
        score, entry_id, sentence = combined_scores[idx]
        if entry_id in seen_entries:
            continue
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute("SELECT validated, metadata, quality_score FROM data_entries WHERE id = ?", (entry_id,))
            validated, metadata, quality_score = c.fetchone()
        if not validated or quality_score < CONFIG['quality_threshold']:
            continue
        doc = nlp(sentence)
        for ent in doc.ents:
            if ent.text.lower() in KNOWLEDGE_BASE:
                score += 0.1
        if intent in ["factoid", "comparison"] and not doc.ents:
            continue
        answers.append((entry_id, sentence, score))
        seen_entries.add(entry_id)
        if len(answers) >= (CONFIG['max_answers'] if intent in ["list", "explanation"] else 1):
            break

    if not answers:
        return "Ei relevanttia tietoa löydetty."

    if intent == "explanation" and len(answers) > 1:
        combined_text = " ".join(sent for _, sent, _ in answers)
        try:
            summary = summarizer.summarize(combined_text, ratio=CONFIG['summary_ratio'])
            result = f"Vastaus:\n{summary or combined_text}"
        except Exception as e:
            logging.warning(f"Yhteenveto epäonnistui: {str(e)}")
            result = f"Vastaus:\n{combined_text}"
    elif intent == "list":
        result = "Vastaukset:\n" + "\n".join(f"- Merkinnästä {entry_id}: {sent}" for entry_id, sent, _ in answers)
    else:
        entry_id, sentence, _ = answers[0]
        result = f"Vastaus (merkinnästä {entry_id}): {sentence}"

    result = clean_answer(result)
    save_faiss_index(index)
    latency = time.time() - start_time
    logging.info(f"Kysymys: {question}, Intentio: {intent}, Viive: {latency:.2f}s, Vastauksia: {len(answers)}")
    redis_client.setex(cache_key, CONFIG['cache_ttl'], result)
    return result

# Tietojen päivitys
def update_data(entry_id, new_content, user_id=None):
    """Päivitä merkinnän sisältö ja upotukset."""
    new_content = normalize_text(new_content)
    is_valid, validation_msg, quality_score = validate_content(new_content)
    with get_db_connection() as conn:
        conn.execute("BEGIN TRANSACTION")
        try:
            c = conn.cursor()
            c.execute("SELECT content FROM data_entries WHERE validated = 1 AND id != ?", (entry_id,))
            existing = [row[0] for row in c.fetchall()]
            if existing:
                new_emb = get_embedding(new_content)
                existing_embs = [get_embedding(e) for e in existing]
                similarities = util.cos_sim(new_emb, existing_embs).numpy().flatten()
                if similarities.max() > 0.9:
                    validation_msg += " Varoitus: Sisältö on hyvin samankaltainen olemassa olevien merkintöjen kanssa."
            c.execute("UPDATE data_entries SET content = ?, validated = ?, quality_score = ? WHERE id = ?",
                      (new_content, 1 if is_valid else 0, quality_score, entry_id))
            c.execute("INSERT INTO history (id, entry_id, content, timestamp) VALUES (?, ?, ?, ?)",
                      (str(uuid.uuid4()), entry_id, new_content, time.strftime("%Y-%m-%d %H:%M:%S")))
            c.execute("DELETE FROM embeddings WHERE entry_id = ?", (entry_id,))
            sentences = [s for s in sent_tokenize(new_content) if s.strip() and len(s) > 5]
            if sentences:
                embeddings = embedder.encode(sentences, batch_size=32, show_progress_bar=False)
                for sent, emb in zip(sentences, embeddings):
                    sent_id = str(uuid.uuid4())
                    compressed_emb = blosc.compress(pickle.dumps(emb), clevel=CONFIG.get('embedding_compression_level', 5))
                    c.execute("INSERT INTO embeddings (sentence_id, entry_id, sentence, embedding) VALUES (?, ?, ?, ?)",
                              (sent_id, entry_id, sent, compressed_emb))
            if safe_es_ping():
                es.update(index='data_entries', id=entry_id, body={
                    'doc': {
                        'content': new_content,
                        'validated': 1 if is_valid else 0,
                        'quality_score': quality_score
                    }
                })
            conn.commit()
        except Exception as e:
            conn.rollback()
            logging.error(f"Tietojen päivitysvirhe: {str(e)}")
            raise e
    clear_cache('list')
    clear_cache('review')
    if user_id:
        log_audit_action(user_id, f"Päivitetty merkintä {entry_id}")
        socketio.emit('collaboration_update', {
            'entry_id': entry_id,
            'action': 'update',
            'user_id': user_id,
            'content_preview': new_content[:100]
        }, room=entry_id)
    return f"Tiedot päivitetty onnistuneesti. {validation_msg} Laatupisteet: {quality_score:.2f}"

# Tietojen poisto
def delete_data(entry_id, user_id=None):
    """Poista merkintä tietokannasta ja Elasticsearchista."""
    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute("DELETE FROM data_entries WHERE id = ?", (entry_id,))
        c.execute("DELETE FROM embeddings WHERE entry_id = ?", (entry_id,))
        if safe_es_ping():
            try:
                es.delete(index='data_entries', id=entry_id)
            except:
                pass
        conn.commit()
    clear_cache('list')
    clear_cache('review')
    if user_id:
        log_audit_action(user_id, f"Poistettu merkintä {entry_id}")
        socketio.emit('collaboration_update', {
            'entry_id': entry_id,
            'action': 'delete',
            'user_id': user_id
        }, room=entry_id)
    return "Tiedot poistettu onnistuneesti."

# Tietojen palautus
def restore_data(entry_id, hist_id=None, user_id=None):
    """Palauta poistettu tai aikaisempi merkinnän versio."""
    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute("SELECT id, content, timestamp FROM history WHERE entry_id = ? ORDER BY timestamp DESC", (entry_id,))
        entries = c.fetchall()
    if not entries:
        return "Ei historiaa löydetty merkinnälle.", []
    if hist_id:
        for entry in entries:
            if entry[0] == hist_id:
                content = normalize_text(entry[1])
                is_valid, validation_msg, quality_score = validate_content(content)
                with get_db_connection() as conn:
                    conn.execute("BEGIN TRANSACTION")
                    try:
                        c = conn.cursor()
                        c.execute("INSERT OR REPLACE INTO data_entries (id, data_type, content, source, validated, metadata, quality_score) VALUES (?, ?, ?, ?, ?, ?, ?)",
                                  (entry_id, 'restored', content, 'history', 1 if is_valid else 0, json.dumps({}), quality_score))
                        c.execute("DELETE FROM embeddings WHERE entry_id = ?", (entry_id,))
                        sentences = [s for s in sent_tokenize(content) if s.strip() and len(s) > 5]
                        if sentences:
                            embeddings = embedder.encode(sentences, batch_size=32, show_progress_bar=False)
                            for sent, emb in zip(sentences, embeddings):
                                sent_id = str(uuid.uuid4())
                                compressed_emb = blosc.compress(pickle.dumps(emb), clevel=CONFIG.get('embedding_compression_level', 5))
                                c.execute("INSERT INTO embeddings (sentence_id, entry_id, sentence, embedding) VALUES (?, ?, ?, ?)",
                                          (sent_id, entry_id, sent, compressed_emb))
                        if safe_es_ping():
                            es.index(index='data_entries', id=entry_id, body={
                                'data_type': 'restored',
                                'content': content,
                                'source': 'history',
                                'validated': 1 if is_valid else 0,
                                'metadata': {},
                                'quality_score': quality_score
                            })
                        conn.commit()
                    except Exception as e:
                        conn.rollback()
                        logging.error(f"Tietojen palautusvirhe: {str(e)}")
                        raise e
                clear_cache('list')
                clear_cache('review')
                if user_id:
                    log_audit_action(user_id, f"Palautettu merkintä {entry_id} historiasta {hist_id}")
                    socketio.emit('collaboration_update', {
                        'entry_id': entry_id,
                        'action': 'restore',
                        'user_id': user_id,
                        'content_preview': content[:100]
                    }, room=entry_id)
                return f"Tiedot palautettu onnistuneesti historiatunnuksesta {hist_id}. {validation_msg} Laatupisteet: {quality_score:.2f}", []
    return "Historiatiedot:", [{'id': e[0], 'timestamp': e[2], 'preview': e[1][:50]} for e in entries]

# Validoimattomien merkintöjen tarkistus
def review_unvalidated(page=1, per_page=20, search=None):
    """Tarkista validoimattomat merkinnät ja anna ehdotuksia."""
    cache_key = f"review:{page}:{per_page}:{search or ''}"
    cached = redis_client.get(cache_key)
    if cached:
        return pickle.loads(cached)
    
    offset = (page - 1) * per_page
    query = "SELECT id, content, source, quality_score FROM data_entries WHERE validated = 0"
    params = []
    if search:
        query += " AND (content LIKE ? OR source LIKE ?)"
        params.extend([f'%{search}%', f'%{search}%'])
    query += " ORDER BY quality_score DESC LIMIT ? OFFSET ?"
    params.extend([per_page, offset])
    
    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute(query, params)
        entries = c.fetchall()
        c.execute("SELECT COUNT(*) FROM data_entries WHERE validated = 0" +
                  (" AND (content LIKE ? OR source LIKE ?)" if search else ""),
                  [p for p in params if p not in [per_page, offset]])
        total = c.fetchone()[0]
    
    if not entries:
        return "Ei validoimattomia merkintöjä.", [], 0
    entries_list = []
    for entry in entries:
        entry_id, content, source, quality_score = entry
        suggestion = "Ei mitään"
        with get_db_connection() as conn:
            c = conn.cursor()
            c.execute("SELECT content FROM data_entries WHERE validated = 1")
            valid_contents = c.fetchall()
        if valid_contents:
            valid_embs = [get_embedding(vc[0]) for vc in valid_contents]
            content_emb = get_embedding(content)
            similarities = util.cos_sim(content_emb, valid_embs).numpy().flatten()
            best_idx = similarities.argmax()
            if similarities[best_idx] > 0.5:
                suggestion = valid_contents[best_idx][0][:100] + "..."
        entries_list.append({
            'id': entry_id,
            'source': source,
            'quality': quality_score,
            'content_preview': content[:100],
            'suggestion': suggestion
        })
    result = ("Validoimattomat merkinnät:", entries_list, total)
    redis_client.setex(cache_key, CONFIG['cache_ttl'], pickle.dumps(result))
    return result

# Tietojen vienti
def export_data(format="json", preview=False):
    """Vie tiedot JSON- tai CSV-muodossa."""
    entries = get_all_data()[0]
    data = [{
        "id": e['id'],
        "data_type": e['type'],
        "content": e['content_preview'] + "...",
        "source": e['source'],
        "validated": e['status'] == "Validoitu",
        "metadata": e['metadata'],
        "quality_score": e['quality'],
        "is_anomaly": e['is_anomaly']
    } for e in entries]
    if preview:
        return json.dumps(data[:5], indent=2)
    if format == "json":
        with open('exported_data.json', 'w') as f:
            json.dump(data, f, indent=2)
        return "Tiedot viety tiedostoon exported_data.json", 'exported_data.json'
    elif format == "csv":
        pd.DataFrame(data).to_csv('exported_data.csv', index=False)
        return "Tiedot viety tiedostoon exported_data.csv", 'exported_data.csv'
    return "Tukematon muoto", None

# Tietojen varmuuskopiointi
def backup_data():
    """Luo tietokannan varmuuskopio."""
    backup_file = f"backup_{time.strftime('%Y%m%d_%H%M%S')}.db"
    with sqlite3.connect('data_store.db') as src, sqlite3.connect(backup_file) as dst:
        src.backup(dst)
    return f"Varmuuskopio luotu: {backup_file}", backup_file

# Analytiikkakaavioiden luonti
def generate_analytics_charts():
    """Luo kaaviot datatyyppien ja laatupisteiden jakautumisesta."""
    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute("SELECT data_type, COUNT(*) FROM data_entries GROUP BY data_type")
        type_counts = c.fetchall()
        c.execute("SELECT quality_score FROM data_entries")
        quality_scores = [row[0] for row in c.fetchall()]
    
    plt.figure(figsize=(8, 6))
    types, counts = zip(*type_counts) if type_counts else ([], [])
    plt.bar(types, counts)
    plt.title('Datatyyppien jakautuminen')
    plt.xlabel('Datatyyppi')
    plt.ylabel('Lukumäärä')
    buf = BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    type_chart = base64.b64encode(buf.getvalue()).decode('utf-8')
    plt.close()

    plt.figure(figsize=(8, 6))
    plt.hist(quality_scores, bins=20, edgecolor='black')
    plt.title('Laatupisteiden jakautuminen')
    plt.xlabel('Laatupisteet')
    plt.ylabel('Taajuus')
    buf = BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    quality_chart = base64.b64encode(buf.getvalue()).decode('utf-8')
    plt.close()

    return type_chart, quality_chart

# Analytiikan haku
def query_analytics():
    """Hae tilastoja kyselyistä ja merkinnöistä."""
    with get_db_connection() as conn:
        c = conn.cursor()
        c.execute("SELECT COUNT(*) FROM embeddings")
        total_sentences = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM data_entries WHERE validated = 1")
        valid_entries = c.fetchone()[0]
        c.execute("SELECT AVG(quality_score) FROM data_entries")
        avg_quality = c.fetchone()[0] or 0.0
        c.execute("SELECT COUNT(*) FROM data_entries WHERE id IN (SELECT id FROM data_entries WHERE validated = 0)")
        anomalies = len(detect_anomalies(get_all_data()[0]))
    intents = {'list': 0, 'definition': 0, 'explanation': 0, 'comparison': 0, 'factoid': 0}
    with open('system.log', 'r') as f:
        questions = 0
        for line in f:
            if "Question:" in line:
                questions += 1
                for intent in intents:
                    if f"Intent: {intent}" in line:
                        intents[intent] += 1
    intent_stats = ", ".join(f"{k}: {v}" for k, v in intents.items())
    type_chart, quality_chart = generate_analytics_charts()
    return {
        'text': f"Analytiikka: {valid_entries} validoitua merkintää, {total_sentences} indeksoitua lausetta, {questions} kysymystä, Keskim. laatu: {avg_quality:.2f}, Poikkeamat: {anomalies}, Intentiotilastot: {intent_stats}",
        'type_chart': type_chart,
        'quality_chart': quality_chart
    }

# Järjestelmän terveystarkistus
def health_check():
    """Tarkista järjestelmän resurssien käyttö ja palveluiden tila."""
    memory = psutil.virtual_memory()
    disk = psutil.disk_usage('/')
    cpu = psutil.cpu_percent()
    redis_status = redis_client.ping()
    es_status = safe_es_ping()
    return f"Muisti: {memory.percent}%, Levy: {disk.percent}%, CPU: {cpu}%, Redis: {'Yhdistetty' if redis_status else 'Ei yhteyttä'}, Elasticsearch: {'Yhdistetty' if es_status else 'Ei yhteyttä'}"

# Ohjeiden näyttäminen
def show_help():
    """Näytä käytettävissä olevat komennot."""
    return """
Komennot:
- upload: Lataa tiedosto (pdf, txt, md, png, jpg, jpeg, docx, csv, zip) tai hakemisto
- question: Kysy kysymys tallennettujen tietojen perusteella
- update: Päivitä merkinnän sisältö
- delete: Poista merkintä
- restore: Palauta poistettu tai aikaisempi merkinnän versio
- list: Listaa kaikki merkinnät
- review: Tarkista validoimattomat merkinnät
- export: Vie tiedot json- tai csv-muodossa
- backup: Luo tietokannan varmuuskopio
- analytics: Näytä kyselytilastot ja kaaviot
- health: Tarkista järjestelmän terveys
- help: Näytä tämä ohje
- exit: Poistu ohjelmasta
"""

# Celery-tehtävät
@celery.task
def process_file_task(file_path, data_type, user_id):
    """Käsittele tiedosto taustalla Celeryllä."""
    processor = FILE_PROCESSORS.get('.' + data_type)
    if not processor:
        socketio.emit('upload_progress', {'message': f"Tukematon tiedostotyyppi: {data_type}", 'user_id': user_id})
        return f"Tukematon tiedostotyyppi: {data_type}"
    content, metadata = processor(file_path)
    if "Virhe" in content:
        socketio.emit('upload_progress', {'message': content, 'user_id': user_id})
        return content
    entry_id, validation_msg, quality_score = store_data(data_type, content, file_path, metadata, user_id)
    result = f"Käsitelty {file_path} tunnuksella: {entry_id}, {validation_msg}, Laatu: {quality_score:.2f}"
    socketio.emit('upload_progress', {'message': result, 'user_id': user_id})
    return result

# Flask-lomakkeet
class UploadForm(FlaskForm):
    file = FileField('Tiedosto tai Zip', validators=[DataRequired()])
    submit = SubmitField('Lataa')

class QuestionForm(FlaskForm):
    question = StringField('Kysymys', validators=[DataRequired()])
    submit = SubmitField('Kysy')

class QueryBuilderForm(FlaskForm):
    content_query = StringField('Sisältöhaku')
    metadata_query = StringField('Metatietohaku')
    data_type = SelectField('Tyyppi', choices=[('', 'Kaikki')] + [(ext[1:], ext[1:].upper()) for ext in CONFIG['supported_extensions'] if ext != '.zip'], coerce=str)
    validated = SelectField('Tila', choices=[('', 'Kaikki'), ('1', 'Validoitu'), ('0', 'Tarvitsee validointia')], coerce=str)
    submit = SubmitField('Hae')

class UpdateForm(FlaskForm):
    entry_id = StringField('Merkinnän tunnus', validators=[DataRequired()])
    content = TextAreaField('Uusi sisältö', validators=[DataRequired()])
    submit = SubmitField('Päivitä')

class DeleteForm(FlaskForm):
    entry_id = StringField('Merkinnän tunnus', validators=[DataRequired()])
    submit = SubmitField('Poista')

class RestoreForm(FlaskForm):
    entry_id = StringField('Merkinnän tunnus', validators=[DataRequired()])
    history_id = StringField('Historiatunnus (valinnainen)')
    submit = SubmitField('Palauta')

class ReviewForm(FlaskForm):
    entry_id = StringField('Merkinnän tunnus', validators=[DataRequired()])
    action = SelectField('Toiminto', choices=[('update', 'Päivitä'), ('delete', 'Poista'), ('skip', 'Ohita')])
    content = TextAreaField('Uusi sisältö (päivitykselle)')
    submit = SubmitField('Käytä')

class ExportForm(FlaskForm):
    format = SelectField('Muoto', choices=[('json', 'JSON'), ('csv', 'CSV')])
    preview = SubmitField('Esikatselu')
    export = SubmitField('Vie')

class LoginForm(FlaskForm):
    email = StringField('Sähköposti', validators=[DataRequired(), Email()])
    password = PasswordField('Salasana', validators=[DataRequired(), Length(min=6)])
    submit = SubmitField('Kirjaudu')

class RegisterForm(FlaskForm):
    email = StringField('Sähköposti', validators=[DataRequired(), Email()])
    password = PasswordField('Salasana', validators=[DataRequired(), Length(min=6)])
    submit = SubmitField('Rekisteröidy')

# Flask-reitit
@app.route('/')
@login_required
def index():
    """Näytä etusivu."""
    return render_template('index.html', health=health_check())

@app.route('/login', methods=['GET', 'POST'])
@limiter.limit("10 per minute")
def login():
    """Käsittele käyttäjän kirjautuminen."""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = LoginForm()
    if form.validate_on_submit():
        with sqlite3.connect('users.db') as conn:
            c = conn.cursor()
            c.execute("SELECT id, email, password FROM users WHERE email = ?", (form.email.data,))
            user = c.fetchone()
        if user and check_password_hash(user[2], form.password.data):
            login_user(User(user[0], user[1], user[2]))
            log_audit_action(user[0], "Kirjauduttu sisään")
            return redirect(url_for('index'))
        flash('Virheellinen sähköposti tai salasana')
    return render_template('login.html', form=form)

@app.route('/register', methods=['GET', 'POST'])
@limiter.limit("5 per minute")
def register():
    """Käsittele käyttäjän rekisteröityminen."""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = RegisterForm()
    if form.validate_on_submit():
        user_id = str(uuid.uuid4())
        hashed_password = generate_password_hash(form.password.data)
        try:
            with sqlite3.connect('users.db') as conn:
                c = conn.cursor()
                c.execute("INSERT INTO users (id, email, password) VALUES (?, ?, ?)",
                          (user_id, form.email.data, hashed_password))
                conn.commit()
            log_audit_action(user_id, "Rekisteröity")
            flash('Rekisteröityminen onnistui! Kirjaudu sisään.')
            return redirect(url_for('login'))
        except sqlite3.IntegrityError:
            flash('Sähköposti on jo rekisteröity.')
    return render_template('register.html', form=form)

@app.route('/logout')
@login_required
def logout():
    """Käsittele käyttäjän uloskirjautuminen."""
    user_id = current_user.id
    logout_user()
    log_audit_action(user_id, "Kirjauduttu ulos")
    return redirect(url_for('login'))

@app.route('/upload', methods=['GET', 'POST'])
@login_required
@limiter.limit("10 per minute")
def upload():
    """Käsittele tiedostojen lataaminen."""
    form = UploadForm()
    if form.validate_on_submit():
        file = form.file.data
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        if filename.endswith('.zip'):
            with tempfile.TemporaryDirectory() as temp_dir:
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                files = [os.path.join(root, f) for root, _, fs in os.walk(temp_dir) for f in fs]
                for full_path in files:
                    file_ext = os.path.splitext(full_path)[1].lower()
                    if file_ext not in CONFIG['supported_extensions']:
                        continue
                    data_type = file_ext[1:]
                    process_file_task.delay(full_path, data_type, current_user.id)
        else:
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext not in CONFIG['supported_extensions']:
                flash(f"Tukematon tiedostotyyppi. Tuetut: {', '.join(CONFIG['supported_extensions'])}")
                return redirect(url_for('upload'))
            data_type = file_ext[1:]
            process_file_task.delay(file_path, data_type, current_user.id)
        flash('Lataus aloitettu. Tarkista edistyminen käyttöliittymästä.')
        return redirect(url_for('upload'))
    return render_template('upload.html', form=form, config=CONFIG)

# SocketIO-yhteyden käsittely
@socketio.on('connect')
def handle_connect():
    if current_user.is_authenticated:
        emit('connected', {'user_id': current_user.id})

# Yhteistyöhuoneeseen liittyminen
@socketio.on('join_collaboration')
def join_collaboration(data):
    entry_id = data.get('entry_id')
    if entry_id:
        join_room(entry_id)
        emit('collaboration_status', {'message': f'Käyttäjä {current_user.id} liittyi yhteistyöhön merkinnällä {entry_id}'}, room=entry_id)

# SocketIO-yhteyden katkaiseminen
@socketio.on('disconnect')
def handle_disconnect():
    if current_user.is_authenticated:
        for room in socketio.server.rooms(current_user.id):
            socketio.leave_room(room)
        logging.info(f"Käyttäjä {current_user.id} katkaisi yhteyden")

@app.route('/question', methods=['GET', 'POST'])
@login_required
@limiter.limit("20 per minute")
def question():
    """Käsittele kysymyksen esittäminen."""
    form = QuestionForm()
    answer = None
    if form.validate_on_submit():
        try:
            answer = answer_question(form.question.data)
            log_audit_action(current_user.id, f"Kysytty kysymys: {form.question.data}")
        except Exception as e:
            flash(f"Virhe kysymykseen vastaamisessa: {str(e)}")
            logging.error(f"Kysymysvirhe: {str(e)}")
    return render_template('question.html', form=form, answer=answer)

@app.route('/query_builder', methods=['GET', 'POST'])
@login_required
def query_builder():
    """Käsittele hakukyselyn rakentaminen."""
    form = QueryBuilderForm()
    page = int(request.args.get('page', 1))
    entries, total = [], 0
    if form.validate_on_submit():
        try:
            _, entries, total = list_entries(
                page,
                CONFIG['items_per_page'],
                form.content_query.data,
                form.data_type.data or None,
                int(form.validated.data) if form.validated.data in ['0', '1'] else None,
                form.metadata_query.data
            )
        except Exception as e:
            flash(f"Virhe kyselyn suorittamisessa: {str(e)}")
            logging.error(f"Kyselyn rakentamisvirhe: {str(e)}")
    return render_template('query_builder.html', form=form, entries=entries, page=page, total=total, per_page=CONFIG['items_per_page'])

@app.route('/list', methods=['GET', 'POST'])
@login_required
def list_data():
    """Listaa tiedot suodattimilla ja sivutuksella."""
    form = QueryBuilderForm()
    page = int(request.args.get('page', 1))
    search = form.content_query.data if form.validate_on_submit() else request.args.get('content_query', '')
    data_type = form.data_type.data if form.validate_on_submit() else request.args.get('data_type', '')
    validated = form.validated.data if form.validate_on_submit() else request.args.get('validated', '')
    metadata_query = form.metadata_query.data if form.validate_on_submit() else request.args.get('metadata_query', '')
    validated = int(validated) if validated in ['0', '1'] else None
    _, entries, total = list_entries(page, CONFIG['items_per_page'], search, data_type or None, validated, metadata_query)
    return render_template('list.html', form=form, entries=entries, page=page, total=total, per_page=CONFIG['items_per_page'], search=search, data_type=data_type, validated=validated, metadata_query=metadata_query)

@app.route('/update', methods=['GET', 'POST'])
@login_required
@limiter.limit("10 per minute")
def update():
    """Käsittele merkinnän päivitys."""
    form = UpdateForm()
    result = None
    if form.validate_on_submit():
        try:
            result = update_data(form.entry_id.data, form.content.data, current_user.id)
            flash(result)
        except Exception as e:
            flash(f"Virhe tietojen päivityksessä: {str(e)}")
            logging.error(f"Päivitysvirhe: {str(e)}")
    return render_template('update.html', form=form, result=result)

@app.route('/delete', methods=['GET', 'POST'])
@login_required
@limiter.limit("10 per minute")
def delete():
    """Käsittele merkinnän poisto."""
    form = DeleteForm()
    result = None
    if form.validate_on_submit():
        try:
            result = delete_data(form.entry_id.data, current_user.id)
            flash(result)
        except Exception as e:
            flash(f"Virhe tietojen poistossa: {str(e)}")
            logging.error(f"Poistovirhe: {str(e)}")
    return render_template('delete.html', form=form, result=result)

@app.route('/restore', methods=['GET', 'POST'])
@login_required
@limiter.limit("10 per minute")
def restore():
    """Käsittele merkinnän palautus."""
    form = RestoreForm()
    result, history = None, []
    if form.validate_on_submit():
        try:
            result, history = restore_data(form.entry_id.data, form.history_id.data or None, current_user.id)
            flash(result)
        except Exception as e:
            flash(f"Virhe tietojen palauttamisessa: {str(e)}")
            logging.error(f"Palautusvirhe: {str(e)}")
    return render_template('restore.html', form=form, result=result, history=history)

@app.route('/review', methods=['GET', 'POST'])
@login_required
@limiter.limit("10 per minute")
def review():
    """Käsittele validoimattomien merkintöjen tarkistus."""
    form = ReviewForm()
    page = int(request.args.get('page', 1))
    search = request.args.get('search', '')
    result, entries, total = review_unvalidated(page, CONFIG['items_per_page'], search)
    if form.validate_on_submit():
        try:
            entry_id = form.entry_id.data
            action = form.action.data
            if action == 'update':
                result = update_data(entry_id, form.content.data, current_user.id)
            elif action == 'delete':
                result = delete_data(entry_id, current_user.id)
            flash(result)
        except Exception as e:
            flash(f"Virhe tietojen tarkistuksessa: {str(e)}")
            logging.error(f"Tarkistusvirhe: {str(e)}")
        return redirect(url_for('review'))
    return render_template('review.html', form=form, entries=entries, page=page, total=total, per_page=CONFIG['items_per_page'], search=search)

@app.route('/export', methods=['GET', 'POST'])
@login_required
@limiter.limit("5 per minute")
def export():
    """Käsittele tietojen vienti."""
    form = ExportForm()
    preview = None
    if form.validate_on_submit():
        try:
            if form.preview.data:
                preview = export_data(form.format.data, preview=True)
            else:
                result, file_path = export_data(form.format.data)
                if file_path:
                    log_audit_action(current_user.id, f"Viety tiedot muodossa {form.format.data}")
                    return send_file(file_path, as_attachment=True)
                flash(result)
        except Exception as e:
            flash(f"Virhe tietojen viennissä: {str(e)}")
            logging.error(f"Vientivirhe: {str(e)}")
    return render_template('export.html', form=form, preview=preview)

@app.route('/backup', methods=['GET', 'POST'])
@login_required
@limiter.limit("5 per minute")
def backup():
    """Käsittele tietokannan varmuuskopiointi."""
    if request.method == 'POST':
        try:
            result, file_path = backup_data()
            log_audit_action(current_user.id, "Luotu varmuuskopio")
            return send_file(file_path, as_attachment=True)
        except Exception as e:
            flash(f"Virhe varmuuskopion luomisessa: {str(e)}")
            logging.error(f"Varmuuskopiovirhe: {str(e)}")
    return render_template('backup.html')

@app.route('/analytics')
@login_required
def analytics():
    """Näytä analytiikkasivu."""
    try:
        analytics = query_analytics()
    except Exception as e:
        flash(f"Virhe analytiikan haussa: {str(e)}")
        logging.error(f"Analytiikkavirhe: {str(e)}")
        analytics = {'text': "Virhe analytiikan haussa.", 'type_chart': None, 'quality_chart': None}
    return render_template('analytics.html', analytics=analytics)

@app.route('/health')
@login_required
def health():
    """Näytä terveystarkistussivu."""
    return render_template('health.html', health=health_check())

@app.route('/help')
@login_required
def help